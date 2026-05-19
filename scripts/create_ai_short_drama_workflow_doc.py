from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "/Users/apple/seedance-ai-video-studio/AI短剧工作流使用教程.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9E2EC", size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("第 ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    p._p.append(fld)
    p.add_run(" 页")
    for r in p.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(100, 116, 139)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("AI 短剧工作流使用教程")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("基于本地 Seedance / Seedream 工作台，从角色、场景、分镜、连续性到剪映拼接的一套完整操作流程。")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(71, 85, 105)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table)
    cells = table.rows[0].cells
    items = [
        ("适用项目", "AI 短剧 / 分镜视频"),
        ("本地地址", "http://localhost:5173/"),
        ("输出目录", "outputs 文件夹"),
    ]
    for cell, (k, v) in zip(cells, items):
        set_cell_shading(cell, "F8FAFC")
        set_cell_border(cell)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(k + "\\n")
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(51, 65, 85)
        r2 = p.add_run(v)
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(15, 23, 42)


def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    return p


def add_h2(doc, text):
    return doc.add_heading(text, level=2)


def add_para(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F1F5F9")
    set_cell_border(cell, "CBD5E1")
    set_cell_margins(cell, 120, 160, 120, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(text.splitlines()):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        r.font.name = "Menlo"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(30, 41, 59)


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "ECFDF5")
    set_cell_border(cell, "A7F3D0")
    set_cell_margins(cell, 120, 160, 120, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title + "：")
    r.bold = True
    r.font.color.rgb = RGBColor(4, 120, 87)
    r.font.size = Pt(10)
    r2 = p.add_run(body)
    r2.font.color.rgb = RGBColor(30, 41, 59)
    r2.font.size = Pt(10)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_workflow_table(doc):
    rows = [
        ("1", "项目设置", "确定项目名、平台、集数、全局风格和负向提示词。"),
        ("2", "角色库", "通过角色描述生成角色草稿和定妆图，确认后保存。"),
        ("3", "场景库", "通过场景描述生成场景草稿和概念图，确认后保存。"),
        ("4", "分镜表", "拆分镜头，设置角色、场景、动作、情绪、台词、比例和时长。"),
        ("5", "连续性", "用上一镜头的结尾动作衔接下一镜头的开头动作。"),
        ("6", "首尾帧", "用首帧控制开始画面，用尾帧控制结束画面。"),
        ("7", "生成视频", "选择参考图，提交 Seedance，自动轮询并下载到 outputs。"),
        ("8", "返修剪辑", "不满意就返修，全部片段导入剪映按编号拼接。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table)
    headers = ["步骤", "模块", "要做什么"]
    widths = [900, 1900, 6560]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "E8EEF5")
        set_cell_border(cell)
        set_cell_margins(cell)
        cell.width = widths[idx]
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        r.bold = True
        r.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.width = widths[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            r.font.size = Pt(10)
            if idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_page_number(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 16, 8),
        ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    add_title(doc)

    add_h1(doc, "1. 这个工作流适合做什么")
    add_para(doc, "这个本地工作台适合一镜一镜地制作 AI 短剧，而不是一次性生成整集。它的核心价值是把角色、场景、分镜、连续性、首尾帧和视频生成串起来，让你像导演一样控制每个镜头。")
    add_workflow_table(doc)

    add_h1(doc, "2. 启动项目")
    add_para(doc, "每次开始制作前，先打开终端进入项目目录并启动服务。")
    add_code(doc, "cd /Users/apple/seedance-ai-video-studio\\nnpm run dev")
    add_para(doc, "看到服务启动后，在浏览器打开：")
    add_code(doc, "http://localhost:5173/")

    add_h1(doc, "3. 设置项目")
    add_bullets(doc, [
        "项目名：短剧名称，例如“雨夜来电”。",
        "平台：抖音短剧、快手短剧、小红书、TikTok 或 YouTube Shorts。",
        "集数：例如“第 1 集”。",
        "全局风格：控制整部剧的影像气质、情绪和构图。",
        "负向提示词：用于避免字幕、水印、卡通风、畸形手指和脸部变形。",
    ])
    add_code(doc, "全局风格示例：\\n现实短剧，电影感，真实人物表演，情绪紧张，竖屏构图，夜雨氛围，冷色调\\n\\n负向提示词示例：\\n不要字幕，不要文字水印，不要卡通风，不要畸形手指，不要脸部变形")

    add_h1(doc, "4. 建立角色库")
    add_para(doc, "角色库是人物一致性的基础。建议先把主要角色都生成定妆图，再开始拆分镜。")
    add_numbered(doc, [
        "在“角色描述”中输入人物设定。",
        "点击“生成角色草稿”。",
        "等待系统生成角色名、身份、角色描述和角色定妆图。",
        "检查并修改角色名、身份、描述。",
        "满意后点击“确认保存角色”。",
    ])
    add_code(doc, "角色描述示例：\\n24岁女主，普通白领，黑色长发，白色针织衫，外表柔弱但很倔强，适合雨夜悬疑短剧")
    add_note(doc, "建议", "主角、男主、反派、关键配角都各自生成一张定妆图。后续分镜可以调用角色图作为参考，提高脸、发型和服装的一致性。")

    add_h1(doc, "5. 建立场景库")
    add_para(doc, "场景库用于保持空间、灯光和氛围一致。常用场景建议先生成概念图。")
    add_numbered(doc, [
        "在“场景描述”中输入环境设定。",
        "点击“生成场景草稿”。",
        "等待系统生成场景名、场景描述和场景概念图。",
        "检查并修改场景名和描述。",
        "满意后点击“确认保存场景”。",
    ])
    add_code(doc, "场景描述示例：\\n雨夜城市街角，便利店门口，玻璃门反射霓虹灯，地面有雨水，夜晚冷色调，悬疑氛围")

    add_h1(doc, "6. 创建分镜")
    add_para(doc, "中间的“分镜表”就是一集短剧的镜头列表。每个镜头尽量只表达一个动作，5 秒左右最稳。")
    add_bullets(doc, [
        "镜头标题：用一句话概括镜头。",
        "角色：选择角色库中保存的角色。",
        "场景：选择场景库中保存的场景。",
        "时长：通常选 5 秒，动作复杂再选 10 秒。",
        "镜别和运镜：例如特写、手持轻微晃动、缓慢推镜。",
        "比例：短剧通常选 9:16。",
        "动作：描述人物正在做什么。",
        "情绪：描述表情和情绪变化。",
        "台词/旁白：需要音频时填写。",
    ])
    add_code(doc, "分镜示例：\\n镜头标题：回头确认\\n角色：林夏\\n场景：雨夜便利店门口\\n时长：5秒\\n镜别：特写\\n运镜：手持轻微晃动\\n比例：9:16\\n动作：她猛地回头看向街角，雨水打湿发梢，身后空无一人\\n情绪：惊慌但强忍镇定\\n台词/旁白：是谁在跟踪我？")

    add_h1(doc, "7. 使用连续性信息")
    add_para(doc, "连续性信息解决的是文字层面的状态衔接。最重要的规则是：上一镜头的结尾动作，要成为下一镜头开头动作的参考。")
    add_code(doc, "上一镜头结尾动作：林夏抬头，身体微微僵住，准备回头\\n下一镜头开头动作：接上一镜头，林夏刚抬头，缓慢回头确认身后")
    add_bullets(doc, [
        "点击“继承上一镜头”可以自动补全当前镜头的连续性信息。",
        "重点检查“开头动作”和“结尾动作”。",
        "场景状态、人物及服装状态、光线氛围要尽量保持稳定。",
        "两个镜头衔接生硬时，可以勾选需要桥接镜头或点击插入桥接镜头。",
    ])

    add_h1(doc, "8. 使用桥接镜头")
    add_para(doc, "桥接镜头的作用不是讲剧情，而是掩盖 AI 视频片段之间的跳跃感。")
    add_bullets(doc, [
        "人物特写切到车门、雨水、手部、玻璃倒影时，可以更自然。",
        "常用类型包括手部特写、眼神特写、呼吸特写、背影、玻璃倒影、衣角/脚步特写、空镜头、环境过渡镜头。",
        "如果两个镜头之间空间、角色或动作变化很大，优先插入桥接镜头。",
    ])

    add_h1(doc, "9. 生成首尾帧")
    add_para(doc, "连续性信息是文字控制，首尾帧是视觉控制。重要镜头建议先生成首帧，再生成视频。")
    add_bullets(doc, [
        "首帧：控制本镜头开始画面，通常会作为视频参考图。",
        "尾帧：控制本镜头结束状态，适合规划下一镜头的衔接。",
        "首帧或尾帧不满意，可以点击返修，写清楚要改哪里。",
    ])
    add_code(doc, "返修意见示例：\\n人物头发要更像角色定妆图，脸不要变，表情更紧张，背景保持便利店门口，不要字幕和文字。")

    add_h1(doc, "10. 生成视频")
    add_bullets(doc, [
        "清晰度：建议 720p。",
        "生成音频：默认开启，如果不需要台词音频可取消。",
        "参考图：默认自动选择最佳参考图，优先使用角色图、场景图、首帧或尾帧。",
        "高级选项：只有临时兜底时才上传本地参考图；本地图片没有公网 URL 时，Ark 可能拒绝。",
        "点击“生成当前镜头”后等待进度条完成。",
    ])
    add_para(doc, "生成成功后，视频会自动保存到本地 outputs 文件夹。")
    add_code(doc, "/Users/apple/seedance-ai-video-studio/outputs")

    add_h1(doc, "11. 视频返修")
    add_para(doc, "如果视频不满意，点击“基于当前视频返修”，写清楚要保留什么、修改什么。")
    add_code(doc, "返修示例：\\n人物脸和角色定妆图不一致，尤其是头发。请保持同一角色，动作仍然是回头看向街角，表情更惊慌，背景保持雨夜便利店门口。")
    add_note(doc, "返修原则", "一次只改一个重点，不要同时要求改脸、改动作、改背景、改运镜。要求越集中，成功率越高。")

    add_h1(doc, "12. 推荐制作顺序")
    add_numbered(doc, [
        "写好项目风格和负向提示词。",
        "生成并保存主要角色定妆图。",
        "生成并保存主要场景概念图。",
        "拆分 5-10 个分镜。",
        "从第 001 镜开始逐个检查连续性。",
        "关键镜头先生成首帧。",
        "满意后生成视频。",
        "不满意就返修。",
        "镜头之间跳跃明显就插入桥接镜头。",
        "全部视频导入剪映拼接。",
    ])

    add_h1(doc, "13. 导入剪映拼接")
    add_bullets(doc, [
        "打开剪映，把 outputs 文件夹里的视频按分镜编号拖入时间线。",
        "删除失败或不满意片段。",
        "按 001、002、003 顺序排列。",
        "桥接镜头放在两个主镜头中间。",
        "添加字幕、音乐、音效和调色。",
        "导出 1080x1920 竖屏视频。",
    ])

    add_h1(doc, "14. 最稳的使用原则")
    add_bullets(doc, [
        "每个镜头只表达一个动作。",
        "角色一致靠角色定妆图、人物及服装状态、首帧和返修提示。",
        "场景一致靠场景概念图、场景状态、光线氛围和首帧。",
        "镜头衔接靠上一镜头结尾动作、下一镜头开头动作、桥接镜头和尾帧。",
        "不要追求一次生成整集，稳定方式是一镜一镜控制。",
    ])

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
